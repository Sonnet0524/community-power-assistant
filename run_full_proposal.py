#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Field Info Agent 完整方案文档生成器
包含专业图表：架构图、流程图、状态图
"""

import subprocess
subprocess.run(['python', '-c', '''
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from PIL import Image, ImageDraw, ImageFont

# 创建临时目录
os.makedirs("temp_images", exist_ok=True)

# 图表生成器
class DiagramGenerator:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.colors = {
            'primary': '#1E3A5F',
            'secondary': '#2E7D32',
            'tertiary': '#E65100',
            'accent': '#00BCD4',
            'text': '#333333'
        }
    
    def get_font(self, size, bold=False):
        try:
            if bold:
                return ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", size)
            return ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", size)
        except:
            return ImageFont.load_default()
    
    def create_architecture_diagram(self):
        width, height = 1600, 1200
        img = Image.new('RGB', (width, height), '#FAFAFA')
        draw = ImageDraw.Draw(img)
        
        title_font = self.get_font(36, True)
        draw.text((width//2, 40), "Field Info Agent 整体业务架构", 
                 fill=self.colors['primary'], font=title_font, anchor="mm")
        
        layers = [
            {'name': '用户层', 'y': 150, 'color': '#E8F5E9', 'border': '#4CAF50',
             'items': ['企业微信APP', '文字/图片消息']},
            {'name': '渠道层', 'y': 330, 'color': '#E3F2FD', 'border': '#2196F3',
             'items': ['WebSocket长连接', '实时双向通信']},
            {'name': '智能层', 'y': 510, 'color': '#FFF3E0', 'border': '#FF9800',
             'items': ['4个核心Skill', '业务逻辑处理']},
            {'name': '工具层', 'y': 690, 'color': '#F3E5F5', 'border': '#9C27B0',
             'items': ['KIMI/PG/MinIO', 'AI/存储/缓存']},
            {'name': '数据层', 'y': 870, 'color': '#ECEFF1', 'border': '#607D8B',
             'items': ['结构化/文件/缓存', '分层存储']}
        ]
        
        box_w = 1200
        box_h = 130
        left_x = 200
        
        for i, layer in enumerate(layers):
            y = layer['y']
            draw.rounded_rectangle([left_x, y, left_x+box_w, y+box_h], 
                                  radius=10, fill=layer['color'], 
                                  outline=layer['border'], width=3)
            
            title_font = self.get_font(20, True)
            draw.text((left_x+20, y+20), layer['name'], 
                     fill=self.colors['primary'], font=title_font)
            
            item_font = self.get_font(15)
            for j, item in enumerate(layer['items']):
                draw.text((left_x+40+j*500, y+75), "• " + item, 
                         fill=self.colors['text'], font=item_font)
            
            if i < len(layers) - 1:
                cx = left_x + box_w // 2
                draw.line([(cx, y+box_h+5), (cx, layers[i+1]['y']-5)], 
                         fill='#999999', width=3)
                draw.polygon([(cx, layers[i+1]['y']-5), 
                             (cx-6, layers[i+1]['y']-15),
                             (cx+6, layers[i+1]['y']-15)], fill='#999999')
        
        filepath = os.path.join(self.output_dir, 'architecture.png')
        img.save(filepath, 'PNG', dpi=(300, 300))
        return filepath
    
    def create_session_state_diagram(self):
        width, height = 1400, 1000
        img = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(img)
        
        title_font = self.get_font(32, True)
        draw.text((width//2, 40), "Session状态流转图", 
                 fill=self.colors['primary'], font=title_font, anchor="mm")
        
        states = [
            {'name': 'IDLE', 'cn': '空闲', 'x': 700, 'y': 180, 'color': '#9E9E9E'},
            {'name': 'PREPARING', 'cn': '准备', 'x': 700, 'y': 340, 'color': '#2196F3'},
            {'name': 'COLLECTING', 'cn': '采集', 'x': 700, 'y': 500, 'color': '#FF9800'},
            {'name': 'ANALYZING', 'cn': '分析', 'x': 700, 'y': 660, 'color': '#9C27B0'},
            {'name': 'COMPLETED', 'cn': '完成', 'x': 700, 'y': 820, 'color': '#4CAF50'}
        ]
        
        for state in states:
            r = 60
            draw.ellipse([state['x']-r, state['y']-r, state['x']+r, state['y']+r],
                        fill=state['color'], outline='white', width=3)
            
            name_font = self.get_font(16, True)
            draw.text((state['x'], state['y']-8), state['name'],
                     fill='white', font=name_font, anchor="mm")
            
            cn_font = self.get_font(12)
            draw.text((state['x'], state['y']+18), state['cn'],
                     fill='white', font=cn_font, anchor="mm")
        
        # 垂直流转箭头
        for i in range(len(states)-1):
            y1 = states[i]['y'] + 65
            y2 = states[i+1]['y'] - 65
            draw.line([(700, y1), (700, y2)], fill='#666666', width=3)
            draw.polygon([(700, y2), (690, y2-10), (710, y2-10)], fill='#666666')
        
        filepath = os.path.join(self.output_dir, 'session_state.png')
        img.save(filepath, 'PNG', dpi=(300, 300))
        return filepath
    
    def create_data_flow_diagram(self):
        width, height = 1600, 900
        img = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(img)
        
        title_font = self.get_font(32, True)
        draw.text((width//2, 35), "核心业务时序图 - 照片分析流程", 
                 fill=self.colors['primary'], font=title_font, anchor="mm")
        
        participants = [
            {'name': '用户', 'x': 150, 'color': '#2196F3'},
            {'name': '企微', 'x': 500, 'color': '#4CAF50'},
            {'name': 'Agent', 'x': 850, 'color': '#FF9800'},
            {'name': 'KIMI', 'x': 1200, 'color': '#9C27B0'}
        ]
        
        for p in participants:
            draw.rectangle([p['x']-70, 90, p['x']+70, 145],
                          fill=p['color'], outline='white', width=2)
            name_font = self.get_font(16, True)
            draw.text((p['x'], 118), p['name'],
                     fill='white', font=name_font, anchor="mm")
            draw.line([(p['x'], 145), (p['x'], 850)], fill='#E0E0E0', width=2)
        
        messages = [
            {'from': 0, 'to': 1, 'y': 200, 'label': '发送照片'},
            {'from': 1, 'to': 2, 'y': 280, 'label': 'WebSocket推送'},
            {'from': 2, 'to': 3, 'y': 360, 'label': '请求AI分析'},
            {'from': 3, 'to': 2, 'y': 440, 'label': '返回结果'},
            {'from': 2, 'to': 1, 'y': 520, 'label': '流式推送'},
            {'from': 1, 'to': 0, 'y': 600, 'label': '展示结果'}
        ]
        
        label_font = self.get_font(14)
        for msg in messages:
            y = msg['y']
            from_x = participants[msg['from']]['x']
            to_x = participants[msg['to']]['x']
            
            draw.line([(from_x, y), (to_x, y)], fill='#666666', width=2)
            
            if to_x > from_x:
                draw.polygon([(to_x, y), (to_x-8, y-5), (to_x-8, y+5)], fill='#666666')
            else:
                draw.polygon([(to_x, y), (to_x+8, y-5), (to_x+8, y+5)], fill='#666666')
            
            mid_x = (from_x + to_x) // 2
            draw.text((mid_x, y-12), msg['label'],
                     fill=self.colors['text'], font=label_font, anchor="mm")
        
        filepath = os.path.join(self.output_dir, 'data_flow.png')
        img.save(filepath, 'PNG', dpi=(300, 300))
        return filepath

# 生成图表
gen = DiagramGenerator("temp_images")
arch_path = gen.create_architecture_diagram()
state_path = gen.create_session_state_diagram()
flow_path = gen.create_data_flow_diagram()

print(f"图表已生成：")
print(f"- 架构图: {arch_path}")
print(f"- 状态图: {state_path}")
print(f"- 流程图: {flow_path}")

# 创建Word文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'SimSun'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# 封面
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('Field Info Agent\\n现场信息收集智能体')
title_run.font.name = 'SimHei'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
title_run.font.size = Pt(26)
title_run.font.bold = True

doc.add_paragraph()

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run('技术方案与业务场景设计书（完整版）')
subtitle_run.font.name = 'SimSun'
subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
subtitle_run.font.size = Pt(16)
subtitle_run.font.bold = True

for _ in range(4):
    doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_para.add_run('版本：V1.0\\n日期：2026年3月\\n编制：PM Agent团队')
info_run.font.name = 'SimSun'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
info_run.font.size = Pt(12)

doc.add_page_break()

# 第一章：项目概述
heading = doc.add_heading(level=1)
run = heading.add_run('一、 项目概述')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(22)
run.font.bold = True

# 插入架构图
doc.add_paragraph()
arch_para = doc.add_paragraph()
arch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = arch_para.add_run()
run.add_picture(arch_path, width=Inches(6))

doc.add_paragraph()
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = caption.add_run('图1-1 Field Info Agent整体业务架构')
cap_run.font.name = 'SimSun'
cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
cap_run.font.size = Pt(10)

doc.add_paragraph()

# 1.1 项目背景
h2 = doc.add_heading(level=2)
run = h2.add_run('1.1 项目背景')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(16)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('随着电力行业数字化转型的深入推进，供电所现场工作人员的信息采集工作面临着诸多挑战。传统的纸质记录、手动整理方式效率低下，信息传递滞后，难以满足现代供电服务的高效、精准要求。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('Field Info Agent基于OpenClaw智能体框架，结合企业微信即时通讯平台和KIMI大模型能力，打造一套专为供电所现场工作设计的信息收集智能体，实现信息采集的智能化、自动化和标准化。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

# 1.2 核心功能表格
doc.add_paragraph()
h2 = doc.add_heading(level=2)
run = h2.add_run('1.2 核心功能')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(16)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('系统围绕供电所现场工作的核心场景，设计了四大核心功能模块：')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

doc.add_paragraph()

# 表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
headers = ['功能模块', '功能描述', '技术特点']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.font.name = 'SimHei'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            r.font.size = Pt(10.5)
            r.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ['驻点工作引导', '提供全流程驻点工作引导，包括工作启动、清单生成、进度跟踪', '自然语言交互、智能意图识别'],
    ['AI照片分析', '对现场照片进行智能识别，分析设备类型、状态、缺陷', 'KIMI多模态、实时分析'],
    ['文档自动生成', '基于采集数据自动生成标准化工作文档', 'Word模板、版本化管理'],
    ['应急处置', '提供应急场景下的快速响应和资源支持', '敏感客户识别、实时指挥']
]

for row_data in rows_data:
    row_cells = table.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text
        for p in row_cells[i].paragraphs:
            for r in p.runs:
                r.font.name = 'SimSun'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                r.font.size = Pt(10.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# 第二章：业务场景设计
heading = doc.add_heading(level=1)
run = heading.add_run('二、 业务场景设计')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(22)

# 插入状态流转图
doc.add_paragraph()
state_para = doc.add_paragraph()
state_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = state_para.add_run()
run.add_picture(state_path, width=Inches(5))

doc.add_paragraph()
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = caption.add_run('图2-1 Session状态流转图')
cap_run.font.name = 'SimSun'
cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
cap_run.font.size = Pt(10)

doc.add_paragraph()

# 场景说明
h2 = doc.add_heading(level=2)
run = h2.add_run('2.1 驻点工作全流程')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(16)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('在Field Info Agent的支持下，驻点工作流程得到全面优化。工作开始前，工作人员只需发送"我今天要去阳光社区驻点"，智能体即可自动识别工作意图，查询历史记录并生成个性化工作清单。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('在现场工作中，工作人员可通过自然语言与智能体交互。到达配电房后发送"开始配电房检查"，智能体将引导完成各项拍摄，AI自动分析设备状态，实时反馈结果。工作完成后，智能体自动生成Word文档并归档。整个过程从传统2小时缩短至30分钟，效率提升75%。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

doc.add_page_break()

# 第三章：技术架构
heading = doc.add_heading(level=1)
run = heading.add_run('三、 技术架构设计')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(22)

# 插入时序图
doc.add_paragraph()
flow_para = doc.add_paragraph()
flow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = flow_para.add_run()
run.add_picture(flow_path, width=Inches(6))

doc.add_paragraph()
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = caption.add_run('图3-1 核心业务时序图 - 照片分析流程')
cap_run.font.name = 'SimSun'
cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
cap_run.font.size = Pt(10)

doc.add_paragraph()

# 技术架构说明
h2 = doc.add_heading(level=2)
run = h2.add_run('3.1 整体架构')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(16)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('Field Info Agent采用分层架构设计，基于OpenClaw智能体框架构建，充分利用企业微信长连接、KIMI大模型多模态能力、PostgreSQL关系型数据库和MinIO对象存储，形成完整的技术体系。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

doc.add_paragraph()

# 架构表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
headers = ['架构层次', '核心组件', '主要功能']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.font.name = 'SimHei'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            r.font.size = Pt(10.5)
            r.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

arch_rows = [
    ['用户交互层', '企业微信APP', '文字、图片、位置消息输入输出'],
    ['渠道接入层', 'WebSocket长连接', '实时双向通信、流式消息推送'],
    ['智能处理层', '4个核心Skill', '业务逻辑处理、意图识别、流程编排'],
    ['工具执行层', '6个工具组件', 'AI分析、数据存储、文档生成'],
    ['数据存储层', 'PG + MinIO + Redis', '结构化数据、文件、会话缓存']
]

for row_data in arch_rows:
    row_cells = table.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text
        for p in row_cells[i].paragraphs:
            for r in p.runs:
                r.font.name = 'SimSun'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                r.font.size = Pt(10.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# 后续章节（简要）
heading = doc.add_heading(level=1)
run = heading.add_run('四、 系统实现方案')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(22)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('系统基于OpenClaw框架开发，采用Node.js技术栈，使用WebSocket长连接模式接入企业微信，内网部署即可，无需公网IP。核心技术包括：KIMI多模态分析、Session状态管理、Word文档生成等。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

doc.add_page_break()

heading = doc.add_heading(level=1)
run = heading.add_run('五、 项目实施计划')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(22)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('项目分为三个阶段：Phase 1 MVP验证（4周）、Phase 2 核心功能开发（4周）、Phase 3 试点验证（4周）。总计12周，资源配置包括2名后端开发、1名DevOps、1名测试、1名产品经理，月成本约600-800元。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

doc.add_page_break()

heading = doc.add_heading(level=1)
run = heading.add_run('六、 总结与展望')
run.font.name = 'SimHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(22)

para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('Field Info Agent将传统耗时2小时的现场记录工作缩短至30分钟，将个人工作经验转化为组织知识资产。未来可扩展设备类型识别、预测性维护、知识图谱等高级功能，持续为电力行业数字化转型贡献力量。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

# 结尾
for _ in range(3):
    doc.add_paragraph()

end_para = doc.add_paragraph()
end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_run = end_para.add_run('— 全文完 —')
end_run.font.name = 'SimSun'
end_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
end_run.font.size = Pt(12)

# 保存
doc.save('Field_Info_Agent_完整方案_含图表_v1.0.docx')
print("✅ Word文档已生成：Field_Info_Agent_完整方案_含图表_v1.0.docx")

# 清理临时文件
import shutil
shutil.rmtree("temp_images", ignore_errors=True)
'''], capture_output=True, text=True, cwd='D:/opencode/github/community-power-assistant')

print("执行结果：", result.stdout)
if result.stderr:
    print("错误：", result.stderr)
