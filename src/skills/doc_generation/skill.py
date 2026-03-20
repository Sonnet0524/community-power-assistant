"""
DocGeneration Skill - 文档自动生成 Skill
支持生成供电简报、应急指引、工作总结三种文档类型
"""

import os
import json
import tempfile
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from openclaw import BaseSkill, SkillContext, SkillResult
from openclaw.tools import KIMITool
from openclaw.tools import MinIOTool


@dataclass
class DocumentSection:
    """文档章节"""
    name: str
    content: str
    level: int = 1


@dataclass
class GeneratedDocument:
    """生成的文档数据"""
    title: str
    doc_type: str
    generated_at: str
    generated_by: str
    sections: Dict[str, str] = field(default_factory=dict)
    photos: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


# 文档模板配置
DOC_TEMPLATES = {
    "briefing": {
        "name": "供电简报",
        "file_name": "{date}_{station}_供电简报.docx",
        "sections": [
            "工作概述",
            "现场情况",
            "设备状态",
            "发现问题",
            "整改建议",
            "附件（照片）"
        ],
        "title_template": "{station}供电简报（{date}）",
        "prompt_template": """
你是一位专业的电力行业文档撰写专家。请根据以下数据生成供电简报内容：

【基本信息】
- 供电所：{station}
- 日期：{date}
- 工作人员：{staff}
- 工作地点：{location}

【采集数据】
{collection_data}

【照片信息】
{photo_count} 张照片已采集

请生成结构化的简报内容，严格按以下JSON格式返回：

{{
    "工作概述": "2-3句话描述本次工作的目的和范围",
    "现场情况": "详细描述现场环境、工作条件、现场布局等",
    "设备状态": "描述所有设备的运行状态、外观情况、参数读数等",
    "发现问题": "列举发现的所有问题，使用编号列表格式",
    "整改建议": "针对发现的问题提出具体的整改措施和建议"
}}

要求：
1. 内容专业、准确、客观
2. 使用规范的电力行业术语
3. 问题描述要具体，包括位置、现象、严重程度
4. 整改建议要有可操作性
"""
    },
    "emergency": {
        "name": "应急指引",
        "file_name": "{date}_{location}_应急指引.docx",
        "sections": [
            "事件概述",
            "影响范围",
            "应急措施",
            "恢复方案",
            "注意事项"
        ],
        "title_template": "{event_type}应急处置指引",
        "prompt_template": """
你是一位资深的电力应急处置专家。请根据以下应急信息生成应急指引：

【事件信息】
- 事件类型：{event_type}
- 发生地点：{location}
- 发生时间：{time}
- 影响范围：{impact}
- 当前状况：{current_status}
- 现场照片：{photo_count} 张

【现场描述】
{scene_description}

请生成专业的应急指引，严格按以下JSON格式返回：

{{
    "事件概述": "简要描述事件经过、原因分析、发展趋势",
    "影响范围": "详细说明受影响的区域、用户数、设备、负荷等",
    "应急措施": "分步骤列出应立即采取的应急措施，包括：1.人员疏散 2.设备隔离 3.备用电源启用 4.现场安全防护",
    "恢复方案": "制定详细的供电恢复方案，包括恢复顺序、操作步骤、预计时间",
    "注意事项": "列出操作过程中的安全注意事项、禁止事项、风险提示"
}}

要求：
1. 内容专业、严谨、可操作性强
2. 强调安全第一原则
3. 措施要具体、可执行
4. 风险提示要全面
"""
    },
    "summary": {
        "name": "工作总结",
        "file_name": "{date}_{period}_工作总结.docx",
        "sections": [
            "工作回顾",
            "完成情况",
            "问题分析",
            "改进措施",
            "下阶段计划"
        ],
        "title_template": "{period}工作总结报告",
        "prompt_template": """
你是一位电力行业运营管理专家。请根据以下数据生成工作总结：

【基本信息】
- 总结周期：{period}
- 所属部门：{department}
- 汇报人：{reporter}

【完成工作】
{completed_tasks}

【发现问题】
{issues}

【统计数据】
- 完成任务数：{task_count}
- 发现问题数：{issue_count}
- 整改完成率：{fix_rate}%
- 客户反馈满意度：{satisfaction}%

请生成完整的工作总结，严格按以下JSON格式返回：

{{
    "工作回顾": "概述本周期主要工作内容、目标达成情况",
    "完成情况": "详细列出完成的任务，包括：1.任务名称 2.完成时间 3.完成质量 4.亮点工作",
    "问题分析": "分析发现的问题，包括：1.问题分类统计 2.问题原因分析 3.改进方向",
    "改进措施": "提出具体的改进措施和优化方案",
    "下阶段计划": "列出下阶段的工作计划、目标、重点任务"
}}

要求：
1. 数据准确、逻辑清晰
2. 分析深入、有数据支撑
3. 改进措施具体可行
4. 计划要有明确的时间节点
"""
    }
}


class DocGenerationSkill(BaseSkill):
    """
    文档自动生成 Skill
    
    根据采集的数据自动生成三种类型的文档：
    1. 供电简报 (briefing) - 日常巡检报告
    2. 应急指引 (emergency) - 应急处置方案
    3. 工作总结 (summary) - 周期性工作总结
    
    特性：
    - 集成 KIMI AI 生成专业内容
    - 自动生成 Word 文档
    - 支持照片嵌入
    - 生成 Markdown 预览
    - 提供分享链接（7天有效期）
    """
    
    NAME = "doc_generation"
    DESCRIPTION = "根据采集数据自动生成供电简报/应急指引/工作总结文档"
    VERSION = "1.0.0"
    
    def __init__(self, config: Optional[Dict] = None):
        """
        初始化 Skill
        
        Args:
            config: 配置字典，包含：
                - kimi_config: KIMI API 配置
                - minio_config: MinIO 配置
                - output_dir: 文档输出目录
        """
        super().__init__(config)
        self.config = config or {}
        
        # 初始化工具
        self.kimi = KIMITool(self.config.get("kimi_config", {}))
        self.minio = MinIOTool(self.config.get("minio_config", {}))
        
        # 文档输出目录
        self.output_dir = Path(self.config.get("output_dir", "/tmp/doc_generation"))
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    async def invoke(self, context: SkillContext) -> SkillResult:
        """
        执行文档生成
        
        Args:
            context: Skill 上下文，包含：
                - doc_type: 文档类型 (briefing/emergency/summary)
                - task_id: 任务ID
                - data: 任务数据字典
                - photos: 照片URL列表
                - user_id: 用户ID（用于记录生成人）
                
        Returns:
            SkillResult: 包含生成的文档URL、预览内容和分享链接
            
        Raises:
            ValueError: 当 doc_type 不支持时
            RuntimeError: 当文档生成失败时
        """
        try:
            # 获取参数
            doc_type = context.params.get("doc_type", "briefing")
            task_id = context.params.get("task_id", "unknown")
            task_data = context.params.get("data", {})
            photos = context.params.get("photos", [])
            user_id = context.params.get("user_id", "system")
            
            # 验证文档类型
            if doc_type not in DOC_TEMPLATES:
                raise ValueError(
                    f"不支持的文档类型: {doc_type}. "
                    f"支持的类型: {list(DOC_TEMPLATES.keys())}"
                )
            
            self.logger.info(f"开始生成文档: type={doc_type}, task_id={task_id}")
            
            # 增强任务数据
            task_data.update({
                "photos": photos,
                "photo_count": len(photos),
                "user_id": user_id
            })
            
            # 使用 KIMI 生成文档内容
            generated_content = await self._generate_content(doc_type, task_data)
            
            # 创建 Word 文档
            doc_url = await self._create_word_document(
                generated_content, 
                doc_type, 
                task_data
            )
            
            # 生成预览
            preview = self._generate_preview(generated_content)
            
            # 创建分享链接
            share_link = await self._create_share_link(doc_url)
            
            self.logger.info(f"文档生成完成: {doc_url}")
            
            return SkillResult(
                response=f"✅ {generated_content.title} 已生成\n\n{preview[:200]}...",
                data={
                    "document_url": doc_url,
                    "preview": preview,
                    "share_link": share_link,
                    "doc_type": doc_type,
                    "generated_at": generated_content.generated_at
                }
            )
            
        except ValueError as e:
            self.logger.error(f"参数错误: {e}")
            return SkillResult(
                response=f"❌ 参数错误: {str(e)}",
                data={"error": str(e)}
            )
        except Exception as e:
            self.logger.error(f"文档生成失败: {e}", exc_info=True)
            return SkillResult(
                response=f"❌ 文档生成失败: {str(e)}",
                data={"error": str(e)}
            )
    
    async def _generate_content(
        self, 
        doc_type: str, 
        task_data: dict
    ) -> GeneratedDocument:
        """
        使用 KIMI 生成文档内容
        
        Args:
            doc_type: 文档类型
            task_data: 任务数据
            
        Returns:
            GeneratedDocument: 生成的文档数据
        """
        template = DOC_TEMPLATES[doc_type]
        
        # 格式化提示词
        prompt = template["prompt_template"].format(**task_data)
        
        # 调用 KIMI 生成
        response = await self.kimi.chat([
            {"role": "system", "content": "你是专业的电力行业文档撰写专家"},
            {"role": "user", "content": prompt}
        ])
        
        # 解析生成的内容
        sections = self._parse_generated_content(response, template["sections"])
        
        # 生成标题
        title = template["title_template"].format(**task_data)
        
        return GeneratedDocument(
            title=title,
            doc_type=doc_type,
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            generated_by=task_data.get("user_id", "system"),
            sections=sections,
            photos=task_data.get("photos", []),
            metadata={
                "template": doc_type,
                "ai_model": "kimi-k2.5"
            }
        )
    
    def _parse_generated_content(
        self, 
        response: str, 
        expected_sections: List[str]
    ) -> Dict[str, str]:
        """
        解析 KIMI 生成的内容
        
        Args:
            response: KIMI 返回的文本
            expected_sections: 期望的章节列表
            
        Returns:
            Dict[str, str]: 章节名到内容的映射
        """
        sections = {}
        
        # 尝试解析 JSON
        try:
            # 提取 JSON 部分
            json_start = response.find("{")
            json_end = response.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                json_str = response[json_start:json_end]
                parsed = json.loads(json_str)
                sections = parsed
        except json.JSONDecodeError:
            self.logger.warning("JSON 解析失败，使用文本解析")
        
        # 如果 JSON 解析失败，使用文本解析
        if not sections:
            current_section = None
            lines = response.split("\n")
            
            for line in lines:
                line = line.strip()
                
                # 识别章节标题
                for section_name in expected_sections:
                    if section_name in line and (":" in line or "：" in line or line.endswith(section_name)):
                        current_section = section_name
                        sections[current_section] = ""
                        break
                else:
                    # 累加内容
                    if current_section and line:
                        sections[current_section] += line + "\n"
        
        # 确保所有章节都有内容
        for section in expected_sections:
            if section not in sections:
                sections[section] = "（待补充）"
        
        return sections
    
    async def _create_word_document(
        self,
        content: GeneratedDocument,
        doc_type: str,
        task_data: dict
    ) -> str:
        """
        创建 Word 文档并上传到 MinIO
        
        Args:
            content: 生成的文档内容
            doc_type: 文档类型
            task_data: 任务数据
            
        Returns:
            str: 文档在 MinIO 的 URL
        """
        template = DOC_TEMPLATES[doc_type]
        
        # 生成文件名
        file_name = template["file_name"].format(
            date=datetime.now().strftime("%Y%m%d"),
            **task_data
        )
        
        # 创建临时文件
        temp_path = self.output_dir / file_name
        
        try:
            # 创建 Word 文档
            doc = Document()
            
            # 设置默认字体
            self._set_default_font(doc)
            
            # 添加标题
            title = doc.add_heading(content.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元信息
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"生成时间：{content.generated_at}").bold = False
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            meta_para2 = doc.add_paragraph()
            meta_para2.add_run(f"生成人：{content.generated_by}").bold = False
            meta_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加分隔线
            doc.add_paragraph("─" * 50)
            
            # 添加各个章节
            for section_name in template["sections"]:
                # 章节标题
                doc.add_heading(section_name, level=1)
                
                # 章节内容
                section_content = content.sections.get(section_name, "")
                if section_content:
                    # 处理多行内容
                    for line in section_content.strip().split("\n"):
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.3)
                
                # 如果是附件章节，添加照片
                if section_name == "附件（照片）" and content.photos:
                    for idx, photo_url in enumerate(content.photos, 1):
                        try:
                            # 下载照片
                            photo_path = await self._download_photo(photo_url, idx)
                            if photo_path and photo_path.exists():
                                # 添加照片说明
                                doc.add_paragraph(f"照片 {idx}")
                                # 添加照片
                                doc.add_picture(str(photo_path), width=Inches(5.0))
                                # 添加空行
                                doc.add_paragraph()
                        except Exception as e:
                            self.logger.warning(f"添加照片失败 {photo_url}: {e}")
                            doc.add_paragraph(f"[照片 {idx} 加载失败]")
            
            # 添加页脚
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            footer = doc.add_paragraph("此文档由 AI 自动生成，请核对后使用")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # 保存文档
            doc.save(str(temp_path))
            
            # 上传到 MinIO
            minio_path = f"generated-docs/{doc_type}/{file_name}"
            doc_url = await self.minio.upload_file(str(temp_path), minio_path)
            
            return doc_url
            
        finally:
            # 清理临时文件
            if temp_path.exists():
                try:
                    temp_path.unlink()
                except Exception:
                    pass
    
    def _set_default_font(self, doc: Document):
        """设置文档默认字体"""
        # 设置中文字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
    
    async def _download_photo(self, photo_url: str, idx: int) -> Optional[Path]:
        """
        下载照片到临时目录
        
        Args:
            photo_url: 照片 URL
            idx: 照片索引
            
        Returns:
            Optional[Path]: 照片本地路径
        """
        try:
            import aiohttp
            
            temp_path = self.output_dir / f"photo_{idx}.jpg"
            
            async with aiohttp.ClientSession() as session:
                async with session.get(photo_url) as response:
                    if response.status == 200:
                        content = await response.read()
                        temp_path.write_bytes(content)
                        return temp_path
                    else:
                        self.logger.warning(f"下载照片失败: {response.status}")
                        return None
                        
        except Exception as e:
            self.logger.warning(f"下载照片异常: {e}")
            return None
    
    def _generate_preview(self, content: GeneratedDocument) -> str:
        """
        生成 Markdown 格式的预览
        
        Args:
            content: 生成的文档内容
            
        Returns:
            str: Markdown 预览文本
        """
        template = DOC_TEMPLATES[content.doc_type]
        
        preview = f"""# {content.title}

**生成时间**：{content.generated_at}  
**生成人**：{content.generated_by}  
**文档类型**：{template['name']}

---

"""
        
        # 添加各个章节预览（限制长度）
        for section_name in template["sections"]:
            section_content = content.sections.get(section_name, "")
            # 限制预览长度
            if len(section_content) > 200:
                section_content = section_content[:200] + "..."
            
            preview += f"## {section_name}\n\n{section_content}\n\n"
        
        # 添加照片信息
        if content.photos:
            preview += f"## 附件\n\n"
            preview += f"共 {len(content.photos)} 张照片\n\n"
        
        preview += "---\n\n*此文档由 AI 自动生成，请核对后使用*"
        
        return preview
    
    async def _create_share_link(self, doc_url: str, expires: int = 604800) -> str:
        """
        生成分享链接
        
        Args:
            doc_url: 文档 URL
            expires: 链接有效期（秒），默认7天
            
        Returns:
            str: 预签名分享链接
        """
        try:
            presigned_url = await self.minio.get_presigned_url(doc_url, expires=expires)
            return presigned_url
        except Exception as e:
            self.logger.error(f"生成分享链接失败: {e}")
            return doc_url  # 返回原始 URL 作为 fallback
    
    async def get_supported_types(self) -> List[Dict]:
        """
        获取支持的文档类型列表
        
        Returns:
            List[Dict]: 文档类型信息列表
        """
        return [
            {
                "type": doc_type,
                "name": info["name"],
                "sections": info["sections"]
            }
            for doc_type, info in DOC_TEMPLATES.items()
        ]
    
    async def validate_params(self, params: dict) -> tuple[bool, str]:
        """
        验证参数是否合法
        
        Args:
            params: 参数字典
            
        Returns:
            tuple[bool, str]: (是否合法, 错误信息)
        """
        required_fields = ["doc_type"]
        
        for field in required_fields:
            if field not in params:
                return False, f"缺少必需参数: {field}"
        
        if params["doc_type"] not in DOC_TEMPLATES:
            return False, f"不支持的文档类型: {params['doc_type']}"
        
        return True, ""
