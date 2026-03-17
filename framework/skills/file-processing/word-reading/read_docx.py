"""
docx 文件阅读理解工具

使用 python-docx 库读取 Word 文件，提取文档内容并返回结构化数据。
"""

from typing import Optional, List, Dict, Any, Union
from pathlib import Path
import json

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    raise ImportError("请安装 python-docx: pip install python-docx")


def get_heading_level(paragraph) -> Optional[int]:
    """
    获取段落的标题级别。
    
    Args:
        paragraph: docx 段落对象
    
    Returns:
        标题级别（1-9），如果不是标题则返回 None
    """
    if paragraph.style.name.startswith('Heading'):
        try:
            level = int(paragraph.style.name.replace('Heading ', ''))
            return level
        except ValueError:
            pass
    return None


def extract_table_data(table) -> Dict[str, Any]:
    """
    提取表格数据。
    
    Args:
        table: docx 表格对象
    
    Returns:
        包含表格数据的字典
    """
    rows_data = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            row_cells.append(cell_text)
        rows_data.append(row_cells)
    
    return {
        "rows": len(table.rows),
        "columns": len(table.columns),
        "data": rows_data
    }


def read_docx(
    file_path: Union[str, Path],
    include_tables: bool = True
) -> Dict[str, Any]:
    """
    读取 docx 文件并返回结构化数据。
    
    Args:
        file_path: Word 文件路径
        include_tables: 是否包含表格内容，默认为 True
    
    Returns:
        包含文件信息的字典，结构如下：
        {
            "file_info": {
                "file_path": "文件路径",
                "paragraphs_count": 段落总数,
                "tables_count": 表格总数
            },
            "title": "文档标题（如果有）",
            "content": [
                {
                    "type": "heading" | "paragraph" | "list",
                    "level": 标题级别（仅标题有）,
                    "text": "文本内容"
                },
                ...
            ],
            "tables": [
                {
                    "rows": 行数,
                    "columns": 列数,
                    "data": [[单元格数据...], ...]
                },
                ...
            ]
        }
    
    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 文件格式错误
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    if file_path.suffix.lower() != '.docx':
        raise ValueError(f"不支持的文件格式: {file_path.suffix}，仅支持 .docx 文件")
    
    try:
        document = Document(str(file_path))
    except Exception as e:
        raise ValueError(f"无法打开文件: {e}")
    
    result: Dict[str, Any] = {
        "file_info": {
            "file_path": str(file_path),
            "paragraphs_count": len(document.paragraphs),
            "tables_count": len(document.tables)
        },
        "title": "",
        "content": [],
        "tables": []
    }
    
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            continue
        
        content_item: Dict[str, Any] = {
            "type": "paragraph",
            "text": text
        }
        
        heading_level = get_heading_level(paragraph)
        if heading_level:
            content_item["type"] = "heading"
            content_item["level"] = heading_level
            
            if not result["title"] and heading_level == 1:
                result["title"] = text
        
        elif paragraph.style and paragraph.style.name.startswith('List'):
            content_item["type"] = "list"
        
        result["content"].append(content_item)
    
    if include_tables and document.tables:
        for table in document.tables:
            table_data = extract_table_data(table)
            result["tables"].append(table_data)
    
    return result


def read_docx_as_markdown(
    file_path: Union[str, Path],
    include_tables: bool = True
) -> str:
    """
    读取 docx 文件并返回 Markdown 格式。
    
    Args:
        file_path: Word 文件路径
        include_tables: 是否包含表格内容
    
    Returns:
        Markdown 格式的字符串
    """
    data = read_docx(file_path, include_tables)
    
    markdown_lines = []
    markdown_lines.append(f"# Word 文件: {data['file_info']['file_path']}\n")
    
    if data["title"]:
        markdown_lines.append(f"**文档标题**: {data['title']}\n")
    
    markdown_lines.append(f"- 段落数: {data['file_info']['paragraphs_count']}")
    markdown_lines.append(f"- 表格数: {data['file_info']['tables_count']}\n")
    
    markdown_lines.append("## 文档内容\n")
    
    for item in data["content"]:
        if item["type"] == "heading":
            level = item.get("level", 1)
            prefix = "#" * (level + 1)
            markdown_lines.append(f"{prefix} {item['text']}\n")
        elif item["type"] == "list":
            markdown_lines.append(f"- {item['text']}")
        else:
            markdown_lines.append(item["text"])
    
    if include_tables and data["tables"]:
        markdown_lines.append("\n## 表格内容\n")
        
        for idx, table in enumerate(data["tables"], 1):
            markdown_lines.append(f"\n### 表格 {idx}\n")
            markdown_lines.append(f"- 行数: {table['rows']}, 列数: {table['columns']}\n")
            
            if table["data"]:
                header = table["data"][0]
                markdown_lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
                markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                
                for row in table["data"][1:]:
                    while len(row) < len(header):
                        row.append("")
                    markdown_lines.append("| " + " | ".join(str(cell) for cell in row[:len(header)]) + " |")
    
    return "\n".join(markdown_lines)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python read_docx.py <文件路径> [--no-tables]")
        print("示例: python read_docx.py document.docx")
        print("      python read_docx.py document.docx --no-tables")
        sys.exit(1)
    
    file_path_arg = sys.argv[1]
    include_tables_arg = "--no-tables" not in sys.argv
    
    try:
        result = read_docx(file_path_arg, include_tables_arg)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
