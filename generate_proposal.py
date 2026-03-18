"""
Field Info Agent - 完整方案文档生成器
基于python-docx创建排版精良的Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_chinese(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    # 根据级别设置字体
    font_sizes = {1: 22, 2: 16, 3: 14, 4: 12}
    font_size = font_sizes.get(level, 12)
    
    set_chinese_font(run, 'SimHei' if level <= 2 else 'SimSun', font_size, bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_chinese(doc, text, bold=False, size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0.5):
    """添加中文段落"""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.first_line_indent = Cm(first_line_indent)
    para.paragraph_format.line_spacing = 1.5
    
    run = para.add_run(text)
    set_chinese_font(run, 'SimSun', size, bold)
    return para

def add_table_chinese(doc, headers, rows):
    """添加中文表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, 'SimHei', 10.5, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, 'SimSun', 10.5)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

def create_field_info_agent_proposal():
    """创建Field Info Agent完整方案文档"""
    
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(10.5)
    
    # ==================== 封面 ====================
    print("创建封面...")
    
    # 空行占位
    for _ in range(6):
        doc.add_paragraph()
    
    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('Field Info Agent\n现场信息收集智能体')
    set_chinese_font(title_run, 'SimHei', 26, bold=True)
    
    doc.add_paragraph()
    
    # 副标题
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('技术方案与业务场景设计书')
    set_chinese_font(subtitle_run, 'SimSun', 16, bold=True)
    
    for _ in range(4):
        doc.add_paragraph()
    
    # 版本信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        '版本：V1.0\n'
        '日期：2026年3月\n'
        '编制：PM Agent团队'
    )
    set_chinese_font(info_run, 'SimSun', 12)
    
    # 分页
    doc.add_page_break()
    
    # ==================== 目录 ====================
    print("创建目录...")
    add_heading_chinese(doc, '目录', level=1)
    
    toc_items = [
        '一、 项目概述',
        '    1.1 项目背景',
        '    1.2 项目目标',
        '    1.3 核心功能',
        '二、 业务场景设计',
        '    2.1 场景一：驻点工作全流程',
        '    2.2 场景二：配电房信息采集',
        '    2.3 场景三：AI照片智能分析',
        '    2.4 场景四：应急处置',
        '三、 技术架构设计',
        '    3.1 整体架构',
        '    3.2 核心组件设计',
        '    3.3 数据存储方案',
        '    3.4 安全设计',
        '四、 系统实现方案',
        '    4.1 开发环境',
        '    4.2 关键技术实现',
        '    4.3 部署方案',
        '五、 项目实施计划',
        '    5.1 项目阶段划分',
        '    5.2 资源需求',
        '    5.3 风险评估',
        '六、 总结与展望'
    ]
    
    for item in toc_items:
        if item.strip().startswith(('一、', '二、', '三、', '四、', '五、', '六、')):
            para = doc.add_paragraph()
            run = para.add_run(item)
            set_chinese_font(run, 'SimHei', 12, bold=True)
        else:
            para = doc.add_paragraph()
            run = para.add_run(item)
            set_chinese_font(run, 'SimSun', 10.5)
            para.paragraph_format.left_indent = Cm(1)
    
    doc.add_page_break()
    
    # ==================== 第一章：项目概述 ====================
    print("创建第一章：项目概述...")
    add_heading_chinese(doc, '一、 项目概述', level=1)
    
    add_heading_chinese(doc, '1.1 项目背景', level=2)
    
    add_paragraph_chinese(doc, 
        '随着电力行业数字化转型的深入推进，供电所现场工作人员的信息采集工作面临着诸多挑战。'
        '传统的纸质记录、手动整理方式效率低下，信息传递滞后，难以满足现代供电服务的高效、精准要求。'
        '现场工作人员在驻点工作中需要同时处理配电房检查、客户走访、应急信息记录等多项任务，'
        '传统的记录方式不仅耗时费力，还容易出现信息遗漏和错误。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '另一方面，随着人工智能技术的快速发展，特别是大语言模型和多模态AI技术的突破，'
        '为现场信息采集工作提供了全新的解决方案。通过智能化的信息收集助手，'
        '可以实现"随手拍、随口说"即可自动完成信息结构化、智能分析和文档生成，'
        '大幅提升工作效率和数据质量。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '本项目旨在基于OpenClaw智能体框架，结合企业微信即时通讯平台和KIMI大模型能力，'
        '打造一套专为供电所现场工作设计的信息收集智能体——Field Info Agent。'
        '该系统将帮助现场工作人员实现信息采集的智能化、自动化和标准化。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '1.2 项目目标', level=2)
    
    add_paragraph_chinese(doc, 
        'Field Info Agent项目的核心目标是构建一个"懂业务、会思考、能生成"的现场工作智能助手，'
        '具体包括以下三个方面：', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '第一，提升信息采集效率。通过自然语言交互和照片智能识别，将传统需要30分钟的记录工作缩短至5分钟，'
        '实现"随手拍、随口说"即可完成信息录入，让现场人员将更多精力投入到实际工作中。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '第二，提高数据质量。通过AI自动分析照片内容、识别设备缺陷、评估运行状态，'
        '减少人为疏漏和错误，确保采集数据的准确性和完整性。同时建立版本化知识库，'
        '实现历史数据的追溯和复用。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '第三，实现知识沉淀。自动生成标准化的工作文档，包括驻点工作总结、设备缺陷报告、'
        '应急通道位置图等，将个人工作经验转化为组织的知识资产，支撑供电所的智慧化运营。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '1.3 核心功能', level=2)
    
    add_paragraph_chinese(doc, 
        'Field Info Agent围绕供电所现场工作的核心场景，设计了四大核心功能模块：', 
        first_line_indent=0.5)
    
    # 核心功能表格
    headers = ['功能模块', '功能描述', '技术特点']
    rows = [
        ['驻点工作引导', '提供全流程的驻点工作引导，包括工作启动、清单生成、进度跟踪', '自然语言交互、智能意图识别、动态清单生成'],
        ['AI照片分析', '对现场照片进行智能识别，分析设备类型、状态、缺陷', 'KIMI多模态、实时分析、批量处理'],
        ['文档自动生成', '基于采集数据自动生成标准化工作文档', 'Word模板、数据自动填充、版本化管理'],
        ['应急处置', '提供应急场景下的快速响应和资源支持', '敏感客户识别、应急方案推送、实时指挥']
    ]
    add_table_chinese(doc, headers, rows)
    
    doc.add_paragraph()
    
    # ==================== 第二章：业务场景设计 ====================
    print("创建第二章：业务场景设计...")
    add_heading_chinese(doc, '二、 业务场景设计', level=1)
    
    add_heading_chinese(doc, '2.1 场景一：驻点工作全流程', level=2)
    
    add_paragraph_chinese(doc, 
        '驻点工作是供电所客户服务的重要组成部分，客户经理需要定期到小区开展配电房检查、'
        '客户走访、应急信息收集等工作。传统模式下，工作人员需要携带纸质记录本，'
        '手动记录各项信息，工作结束后再整理成电子文档，整个过程耗时且容易遗漏。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '在Field Info Agent的支持下，驻点工作流程得到全面优化。工作开始前，工作人员只需在企业微信中'
        '发送"我今天要去阳光社区驻点"，智能体即可自动识别工作意图，查询该社区的历史记录，'
        '包括上次驻点时间、重点客户名单、遗留问题清单等，并生成个性化的工作清单。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '在现场工作中，工作人员可以通过自然语言与智能体交互。例如，到达配电房后发送"开始配电房检查"，'
        '智能体将引导完成各项检查项的拍摄，包括变压器整体、高压侧、低压侧、环境照片等。'
        '每张照片上传后，AI会自动分析设备状态和潜在缺陷，实时反馈分析结果。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '工作完成后，智能体自动汇总所有采集信息，生成包含照片、分析结果、工作记录的Word文档，'
        '并归档到知识库中供团队成员查阅。整个过程从传统的2小时缩短至30分钟，效率提升75%。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '2.2 场景二：配电房信息采集', level=2)
    
    add_paragraph_chinese(doc, 
        '配电房是小区供电系统的核心设施，定期检查配电房设备状态是保障供电安全的关键工作。'
        '传统方式下，工作人员需要手动记录变压器型号、容量、运行状态等信息，并拍摄照片作为佐证，'
        '回到办公室后再整理成电子档案。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'Field Info Agent通过AI视觉技术革新了配电房信息采集方式。工作人员只需拍摄变压器铭牌照片，'
        'KIMI多模态模型即可自动识别设备型号、容量、厂家、出厂日期等关键信息，准确率达到95%以上。'
        '对于设备整体照片，AI能够评估设备外观状态，识别锈蚀、漏油、松动等缺陷，并给出处理建议。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '系统还支持批量照片分析功能。当工作人员完成配电房所有设备拍摄后，可以一键触发批量分析，'
        'AI将在后台对多张照片进行综合分析，识别设备间的关联问题，生成整体的配电房健康评估报告。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '2.3 场景三：AI照片智能分析', level=2)
    
    add_paragraph_chinese(doc, 
        '照片是现场工作最真实的第一手资料，传统方式下照片只是作为存档，价值未得到充分挖掘。'
        'Field Info Agent通过KIMI大模型的多模态能力，赋予照片全新的价值。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '系统支持设备类型自动识别，包括变压器、高压柜、低压柜、电缆、计量装置等常见电力设备。'
        '在识别设备类型的基础上，AI会对设备运行状态进行评估，分为正常、注意、异常、危险四个等级。'
        '对于发现的缺陷，系统会标注缺陷类别（外观、连接、绝缘、运行等）、严重程度和建议处理措施。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '特别值得一提的是安全隐患识别能力。AI能够从照片中识别出变压器漏油、设备锈蚀、电缆老化、'
        '标识缺失等安全隐患，及时提醒工作人员关注，将事故隐患消除在萌芽状态。'
        '这一功能对于保障供电安全和人身安全具有重要意义。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '2.4 场景四：应急处置', level=2)
    
    add_paragraph_chinese(doc, 
        '电力应急抢修是供电所工作的重要组成部分，需要快速响应、准确判断、有效处置。'
        'Field Info Agent为应急处置提供全方位的智能支持。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '当发生停电故障时，抢修人员只需发送"阳光小区停电"，智能体即可自动启动应急模式，'
        '快速查询小区的应急接入点位置、电缆型号、所需长度等关键信息，并推送应急发电车接入指引。'
        '同时，系统会自动识别受影响范围内的敏感客户，如独居老人、医院、学校等，生成关怀清单，'
        '提醒抢修人员优先保障。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '在抢修过程中，智能体支持实时记录处理进展，包括到达现场时间、故障原因、处理措施、恢复时间等，'
        '自动生成应急处理报告，为后续的事故分析和经验总结提供数据支撑。', 
        first_line_indent=0.5)
    
    doc.add_page_break()
    
    # ==================== 第三章：技术架构设计 ====================
    print("创建第三章：技术架构设计...")
    add_heading_chinese(doc, '三、 技术架构设计', level=1)
    
    add_heading_chinese(doc, '3.1 整体架构', level=2)
    
    add_paragraph_chinese(doc, 
        'Field Info Agent采用分层架构设计，基于OpenClaw智能体框架构建，'
        '充分利用企业微信长连接、KIMI大模型多模态能力、PostgreSQL关系型数据库和MinIO对象存储，'
        '形成完整的技术体系。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '系统架构分为五个层次：用户交互层、渠道接入层、智能处理层、工具执行层和数据存储层。'
        '用户交互层通过企业微信APP提供自然、便捷的交互体验；渠道接入层基于WebSocket长连接'
        '实现实时双向通信；智能处理层包含四大核心Skill，实现业务逻辑；工具执行层提供具体的'
        '技术服务能力；数据存储层负责数据的持久化和版本管理。', 
        first_line_indent=0.5)
    
    # 架构特点表格
    headers = ['架构层次', '核心组件', '主要功能']
    rows = [
        ['用户交互层', '企业微信APP', '提供文字、图片、位置消息的输入输出'],
        ['渠道接入层', 'WebSocket长连接', '实时双向通信、流式消息推送'],
        ['智能处理层', '4个核心Skill', '业务逻辑处理、意图识别、流程编排'],
        ['工具执行层', '6个工具组件', 'AI分析、数据存储、文档生成'],
        ['数据存储层', 'PG + MinIO + Redis', '结构化数据、文件、会话缓存']
    ]
    add_table_chinese(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_chinese(doc, '3.2 核心组件设计', level=2)
    
    add_paragraph_chinese(doc, 
        '智能处理层是系统的核心，包含四个精心设计的Skill，分别负责不同的业务场景。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'StationWorkGuide（驻点工作引导）是主控Skill，负责整个驻点工作流程的引导和协调。'
        '它通过自然语言理解用户意图，自动识别工作阶段，动态生成工作清单，引导现场信息采集流程。'
        '支持的工作类型包括配电房检查、客户走访、应急通道采集等。核心能力包括零命令自然语言交互、'
        '上下文感知状态管理、动态工作清单生成和多轮对话引导。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'VisionAnalysis（AI照片分析）负责照片的智能化处理。它利用KIMI 2.5的多模态能力，'
        '对现场照片进行设备识别、缺陷检测、状态评估等分析。支持单图实时分析和批量深度分析两种模式，'
        '分析结果自动保存到数据库，为后续文档生成提供数据支撑。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'DocGeneration（文档自动生成）根据采集的数据自动生成标准化的工作文档。'
        '支持驻点工作记录表、设备缺陷报告、安全隐患整改通知单、应急通道位置图等多种文档类型。'
        '采用模板化设计，确保文档格式统一、内容完整，并支持版本化管理。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'EmergencyGuide（应急处理指引）专门针对应急场景设计，提供快速响应和资源支持。'
        '当检测到应急关键词时自动触发，能够快速查询应急资料、分析影响范围、识别敏感客户、'
        '推送应急方案，并支持处理过程的实时记录。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '3.3 数据存储方案', level=2)
    
    add_paragraph_chinese(doc, 
        '系统采用分层存储策略，根据数据类型选择合适的存储方案，确保性能、成本和可靠性的平衡。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'PostgreSQL关系型数据库存储结构化业务数据，包括会话信息、采集记录、照片分析结果、'
        '生成文档信息等。数据库采用版本化设计，所有修改都创建新版本，历史记录永不删除，'
        '支持完整的数据追溯。关键表包括field_sessions（会话表）、field_collections（采集记录表）、'
        'photo_analysis（照片分析表）、generated_documents（生成文档表）等。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'MinIO对象存储负责非结构化数据的存储，包括原始照片、生成的Word文档、文档模板等。'
        '采用桶隔离策略，按业务类型划分不同的存储桶，如field-photos存储照片、'
        'field-documents存储文档。支持版本控制，所有文件的历史版本都可追溯。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'Redis内存数据库用于会话状态缓存和并发控制。活跃会话存储在Redis中，'
        '设置7200秒（2小时）的过期时间，无操作自动清理。通过分布式锁机制防止并发冲突，'
        '确保同一用户的多个请求能够正确处理。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '3.4 安全设计', level=2)
    
    add_paragraph_chinese(doc, 
        '安全是系统设计的首要原则。Field Info Agent从多个维度构建了完整的安全防护体系。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '数据安全方面，所有敏感数据都进行加密存储，包括数据库连接密码、API密钥、用户个人信息等。'
        '照片和文档存储在本地MinIO中，不依赖第三方云存储，确保数据主权。'
        'WebSocket长连接基于wss协议，全程加密传输，无需额外处理消息解密。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '访问控制方面，系统实现基于角色的权限管理（RBAC）。用户只能访问本供电所的数据，'
        '管理员可以查看所有数据。所有操作都记录审计日志，包括操作人、时间、内容，'
        '保留365天以备追溯。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '网络安全方面，系统部署在企业内网，通过WebSocket长连接主动向外建连，'
        '无需开放公网端口，有效避免外部攻击。数据库和存储服务都运行在内网，'
        '与外部网络物理隔离，符合电力行业安全规范。', 
        first_line_indent=0.5)
    
    doc.add_page_break()
    
    # ==================== 第四章：系统实现方案 ====================
    print("创建第四章：系统实现方案...")
    add_heading_chinese(doc, '四、 系统实现方案', level=1)
    
    add_heading_chinese(doc, '4.1 开发环境', level=2)
    
    add_paragraph_chinese(doc, 
        'Field Info Agent基于OpenClaw框架开发，采用Node.js技术栈。'
        '开发环境需要Node.js 18+、TypeScript 5+、OpenClaw SDK。', 
        first_line_indent=0.5)
    
    # 技术栈表格
    headers = ['类别', '技术选型', '版本要求']
    rows = [
        ['开发框架', 'OpenClaw', '1.0+'],
        ['编程语言', 'TypeScript', '5.0+'],
        ['运行时', 'Node.js', '18+'],
        ['数据库', 'PostgreSQL', '14+'],
        ['对象存储', 'MinIO', '最新版'],
        ['缓存', 'Redis', '7.0+'],
        ['AI模型', 'KIMI K2.5', 'API接入'],
        ['容器化', 'Docker', '20+'],
        ['编排', 'Docker Compose', '2.0+']
    ]
    add_table_chinese(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_chinese(doc, '4.2 关键技术实现', level=2)
    
    add_paragraph_chinese(doc, 
        'WebSocket长连接接入是系统的关键技术点。与企业微信建立长连接需要完成三个步骤：'
        '首先建立WebSocket连接到wss://openws.work.weixin.qq.com，然后发送aibot_subscribe订阅请求'
        '进行身份验证，最后保持心跳保活（建议30秒间隔）。OpenClaw框架封装了这些底层细节，'
        '开发者只需配置BotID和Secret即可。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'KIMI多模态分析是另一个核心技术。系统利用KIMI 2.5的图像理解能力，'
        '对电力设备照片进行专业分析。通过精心设计的Prompt模板，引导AI输出结构化的JSON结果，'
        '包括设备类型、运行状态、缺陷列表、安全建议等。对于批量照片，采用分批处理策略，'
        '每批10张照片，支持进度实时反馈。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'Session状态管理采用Redis存储活跃会话，PostgreSQL持久化历史记录。'
        '会话状态包括准备阶段、采集阶段、分析阶段、完成阶段四个状态，'
        '每个阶段都有详细的上下文信息。通过事件溯源模式记录所有操作日志，'
        '支持任意时间点的状态重建。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        'Word文档生成使用模板化技术。系统预置多种文档模板，包括驻点工作记录表、'
        '设备缺陷报告、应急通道位置图等。生成时动态填充数据、插入照片、渲染表格，'
        '最终输出标准格式的docx文件。生成的文档自动上传到MinIO并记录到数据库。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '4.3 部署方案', level=2)
    
    add_paragraph_chinese(doc, 
        '系统采用Docker容器化部署，使用Docker Compose进行服务编排。'
        '整个系统包含五个服务：OpenClaw Agent应用、PostgreSQL数据库、MinIO对象存储、'
        'Redis缓存和一个Nginx反向代理。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '推荐的部署架构是单机部署（4核8G配置），所有服务运行在同一台服务器上，'
        '简化运维复杂度。对于有高可用要求的场景，可以采用主从部署架构，'
        '数据库和存储配置主从复制，应用层通过负载均衡实现故障转移。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '由于采用WebSocket长连接模式，服务器无需公网IP，部署在企业内网即可。'
        '只需确保服务器能够访问企业微信的WebSocket网关（wss://openws.work.weixin.qq.com），'
        '防火墙放通443端口的出站连接。这种部署方式大大降低了网络配置复杂度，'
        '同时提升了系统安全性。', 
        first_line_indent=0.5)
    
    doc.add_page_break()
    
    # ==================== 第五章：项目实施计划 ====================
    print("创建第五章：项目实施计划...")
    add_heading_chinese(doc, '五、 项目实施计划', level=1)
    
    add_heading_chinese(doc, '5.1 项目阶段划分', level=2)
    
    add_paragraph_chinese(doc, 
        '项目采用敏捷开发模式，分为三个阶段，总计12周。每个阶段都有明确的交付物和验收标准。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '第一阶段（第1-4周）：MVP验证。完成基础环境搭建、企业微信长连接接入、'
        'KIMI多模态能力验证、简单驻点工作流开发。交付可运行的Demo版本，'
        '验证核心流程的技术可行性。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '第二阶段（第5-8周）：核心功能开发。完成四大Skill的全面开发，'
        '包括驻点工作引导、AI照片分析、文档自动生成、应急处置指引。'
        '实现完整的业务闭环，交付功能完整的测试版本。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '第三阶段（第9-12周）：试点验证。选择试点供电所进行实际部署，'
        '培训现场人员使用，收集反馈意见，持续优化系统。完成性能调优、'
        'Bug修复、文档完善，交付生产版本和试点报告。', 
        first_line_indent=0.5)
    
    # 项目阶段表格
    headers = ['阶段', '周期', '主要任务', '交付物']
    rows = [
        ['Phase 1: MVP验证', '第1-4周', '环境搭建、基础接入、简单工作流', '可运行Demo'],
        ['Phase 2: 核心功能', '第5-8周', '四大Skill全面开发', '测试版本'],
        ['Phase 3: 试点验证', '第9-12周', '试点部署、培训、优化', '生产版本+报告']
    ]
    add_table_chinese(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_chinese(doc, '5.2 资源需求', level=2)
    
    add_paragraph_chinese(doc, 
        '项目资源包括人力资源和基础设施两部分。人力资源方面，'
        '建议配置2名后端开发工程师（负责OpenClaw开发）、1名DevOps工程师（负责部署运维）、'
        '1名测试工程师（负责功能测试）、1名产品经理（负责需求协调）。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '基础设施方面，推荐配置如下：云服务器4核8G（月费用约400元）、'
        'KIMI K2.5 API调用（按量计费，月费用约200-400元）、PostgreSQL和MinIO本地部署（免费）、'
        '总计月成本约600-800元，远低于传统方案的数千元。', 
        first_line_indent=0.5)
    
    add_heading_chinese(doc, '5.3 风险评估', level=2)
    
    add_paragraph_chinese(doc, 
        '项目面临的风险主要包括技术风险、业务风险和实施风险三类。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '技术风险方面，主要关注KIMI多模态模型的准确率。虽然KIMI 2.5在图像理解方面表现优秀，'
        '但电力设备的专业性较强，可能需要针对特定场景进行Prompt优化。'
        '应对措施包括建立人工审核机制、持续优化Prompt模板、准备备选AI模型。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '业务风险方面，主要关注现场用户的接受度。部分老员工可能不习惯使用智能工具，'
        '存在抵触情绪。应对措施包括充分的培训、简化操作流程、提供现场技术支持、'
        '建立激励机制鼓励使用。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '实施风险方面，主要关注API权限申请和试点推进。企业微信和KIMI API的审批可能需要时间，'
        '试点供电所的选择和协调也需要提前沟通。应对措施包括提前申请API权限、'
        '尽早确定试点单位、制定详细的实施计划。', 
        first_line_indent=0.5)
    
    doc.add_page_break()
    
    # ==================== 第六章：总结与展望 ====================
    print("创建第六章：总结与展望...")
    add_heading_chinese(doc, '六、 总结与展望', level=1)
    
    add_paragraph_chinese(doc, 
        'Field Info Agent是一个面向供电所现场工作场景的智能信息收集系统，'
        '基于OpenClaw智能体框架，充分利用企业微信长连接、KIMI大模型多模态能力等先进技术，'
        '为现场工作人员提供智能化的信息采集、分析和文档生成服务。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '项目的核心价值在于将传统耗时2小时的现场记录工作缩短至30分钟，'
        '将个人工作经验转化为组织的知识资产，实现供电所现场工作的数字化、智能化转型。'
        '通过AI视觉分析，系统能够自动识别设备缺陷和安全隐患，将被动响应转变为主动预防，'
        '对于提升供电安全和服务质量具有重要意义。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '展望未来，Field Info Agent可以进一步扩展功能边界。一是扩展设备类型识别范围，'
        '覆盖更多电力设备；二是引入预测性维护能力，基于历史数据预测设备故障；'
        '三是构建供电所知识图谱，实现知识的智能检索和推荐；四是接入更多数据源，'
        '如IoT传感器数据，实现多维度的设备健康评估。', 
        first_line_indent=0.5)
    
    add_paragraph_chinese(doc, 
        '我们相信，随着人工智能技术的不断发展和应用的不断深入，'
        'Field Info Agent将成为供电所现场工作的得力助手，'
        '为电力行业的数字化转型贡献力量。', 
        first_line_indent=0.5)
    
    # 添加结尾空行
    for _ in range(3):
        doc.add_paragraph()
    
    # 结尾标注
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_para.add_run('— 全文完 —')
    set_chinese_font(end_run, 'SimSun', 12)
    
    # 保存文档
    output_path = 'D:\\opencode\\github\\community-power-assistant/Field_Info_Agent_技术方案_v1.0.docx'
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")
    
    return output_path

if __name__ == '__main__':
    create_field_info_agent_proposal()
