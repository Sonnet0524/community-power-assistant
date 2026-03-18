#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Field Info Agent 完整方案文档生成器（含图表）
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import shutil
from PIL import Image, ImageDraw, ImageFont

# 创建临时目录
os.makedirs("temp_images", exist_ok=True)

class DiagramGenerator:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.colors = {
            'primary': '#1E3A5F',
            'secondary': '#2E7D32',
            'tertiary': '#E65100',
            'accent': '#00BCD4',
            'text': '#333333'
        }
    
    def get_font(self, size, bold=False):
        try:
            if bold:
                return ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", size)
            return ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", size)
        except:
            try:
                return ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", size)
            except:
                return ImageFont.load_default()
    
    def create_architecture_diagram(self):
        width, height = 1600, 1200
        img = Image.new('RGB', (width, height), '#FAFAFA')
        draw = ImageDraw.Draw(img)
        
        title_font = self.get_font(36, True)
        draw.text((width//2, 40), "Field Info Agent 整体业务架构", 
                 fill=self.colors['primary'], font=title_font, anchor="mm")
        
        layers = [
            {'name': '用户层', 'y': 150, 'color': '#E8F5E9', 'border': '#4CAF50',
             'items': ['企业微信APP', '文字/图片消息']},
            {'name': '渠道层', 'y': 330, 'color': '#E3F2FD', 'border': '#2196F3',
             'items': ['WebSocket长连接', '实时双向通信']},
            {'name': '智能层', 'y': 510, 'color': '#FFF3E0', 'border': '#FF9800',
             'items': ['4个核心Skill', '业务逻辑处理']},
            {'name': '工具层', 'y': 690, 'color': '#F3E5F5', 'border': '#9C27B0',
             'items': ['KIMI/PG/MinIO', 'AI/存储/缓存']},
            {'name': '数据层', 'y': 870, 'color': '#ECEFF1', 'border': '#607D8B',
             'items': ['结构化/文件/缓存', '分层存储']}
        ]
        
        box_w = 1200
        box_h = 130
        left_x = 200
        
        for i, layer in enumerate(layers):
            y = layer['y']
            draw.rounded_rectangle([left_x, y, left_x+box_w, y+box_h], 
                                  radius=10, fill=layer['color'], 
                                  outline=layer['border'], width=3)
            
            title_font = self.get_font(20, True)
            draw.text((left_x+20, y+20), layer['name'], 
                     fill=self.colors['primary'], font=title_font)
            
            item_font = self.get_font(15)
            for j, item in enumerate(layer['items']):
                draw.text((left_x+40+j*500, y+75), "• " + item, 
                         fill=self.colors['text'], font=item_font)
            
            if i < len(layers) - 1:
                cx = left_x + box_w // 2
                draw.line([(cx, y+box_h+5), (cx, layers[i+1]['y']-5)], 
                         fill='#999999', width=3)
                draw.polygon([(cx, layers[i+1]['y']-5), 
                             (cx-6, layers[i+1]['y']-15),
                             (cx+6, layers[i+1]['y']-15)], fill='#999999')
        
        filepath = os.path.join(self.output_dir, 'architecture.png')
        img.save(filepath, 'PNG', dpi=(300, 300))
        return filepath
    
    def create_session_state_diagram(self):
        width, height = 1400, 1000
        img = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(img)
        
        title_font = self.get_font(32, True)
        draw.text((width//2, 40), "Session状态流转图", 
                 fill=self.colors['primary'], font=title_font, anchor="mm")
        
        states = [
            {'name': 'IDLE', 'cn': '空闲', 'x': 700, 'y': 180, 'color': '#9E9E9E'},
            {'name': 'PREPARING', 'cn': '准备', 'x': 700, 'y': 340, 'color': '#2196F3'},
            {'name': 'COLLECTING', 'cn': '采集', 'x': 700, 'y': 500, 'color': '#FF9800'},
            {'name': 'ANALYZING', 'cn': '分析', 'x': 700, 'y': 660, 'color': '#9C27B0'},
            {'name': 'COMPLETED', 'cn': '完成', 'x': 700, 'y': 820, 'color': '#4CAF50'}
        ]
        
        for state in states:
            r = 60
            draw.ellipse([state['x']-r, state['y']-r, state['x']+r, state['y']+r],
                        fill=state['color'], outline='white', width=3)
            
            name_font = self.get_font(16, True)
            draw.text((state['x'], state['y']-8), state['name'],
                     fill='white', font=name_font, anchor="mm")
            
            cn_font = self.get_font(12)
            draw.text((state['x'], state['y']+18), state['cn'],
                     fill='white', font=cn_font, anchor="mm")
        
        for i in range(len(states)-1):
            y1 = states[i]['y'] + 65
            y2 = states[i+1]['y'] - 65
            draw.line([(700, y1), (700, y2)], fill='#666666', width=3)
            draw.polygon([(700, y2), (690, y2-10), (710, y2-10)], fill='#666666')
        
        filepath = os.path.join(self.output_dir, 'session_state.png')
        img.save(filepath, 'PNG', dpi=(300, 300))
        return filepath
    
    def create_data_flow_diagram(self):
        width, height = 1600, 900
        img = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(img)
        
        title_font = self.get_font(32, True)
        draw.text((width//2, 35), "核心业务时序图", 
                 fill=self.colors['primary'], font=title_font, anchor="mm")
        
        participants = [
            {'name': '用户', 'x': 150, 'color': '#2196F3'},
            {'name': '企微', 'x': 500, 'color': '#4CAF50'},
            {'name': 'Agent', 'x': 850, 'color': '#FF9800'},
            {'name': 'KIMI', 'x': 1200, 'color': '#9C27B0'}
        ]
        
        for p in participants:
            draw.rectangle([p['x']-70, 90, p['x']+70, 145],
                          fill=p['color'], outline='white', width=2)
            name_font = self.get_font(16, True)
            draw.text((p['x'], 118), p['name'],
                     fill='white', font=name_font, anchor="mm")
            draw.line([(p['x'], 145), (p['x'], 850)], fill='#E0E0E0', width=2)
        
        messages = [
            {'from': 0, 'to': 1, 'y': 200, 'label': '发送照片'},
            {'from': 1, 'to': 2, 'y': 280, 'label': 'WebSocket推送'},
            {'from': 2, 'to': 3, 'y': 360, 'label': '请求AI分析'},
            {'from': 3, 'to': 2, 'y': 440, 'label': '返回结果'},
            {'from': 2, 'to': 1, 'y': 520, 'label': '流式推送'},
            {'from': 1, 'to': 0, 'y': 600, 'label': '展示结果'}
        ]
        
        label_font = self.get_font(14)
        for msg in messages:
            y = msg['y']
            from_x = participants[msg['from']]['x']
            to_x = participants[msg['to']]['x']
            
            draw.line([(from_x, y), (to_x, y)], fill='#666666', width=2)
            
            if to_x > from_x:
                draw.polygon([(to_x, y), (to_x-8, y-5), (to_x-8, y+5)], fill='#666666')
            else:
                draw.polygon([(to_x, y), (to_x+8, y-5), (to_x+8, y+5)], fill='#666666')
            
            mid_x = (from_x + to_x) // 2
            draw.text((mid_x, y-12), msg['label'],
                     fill=self.colors['text'], font=label_font, anchor="mm")
        
        filepath = os.path.join(self.output_dir, 'data_flow.png')
        img.save(filepath, 'PNG', dpi=(300, 300))
        return filepath

# 生成图表
print("正在生成图表...")
gen = DiagramGenerator("temp_images")
arch_path = gen.create_architecture_diagram()
print(f"✓ 架构图: {arch_path}")
state_path = gen.create_session_state_diagram()
print(f"✓ 状态图: {state_path}")
flow_path = gen.create_data_flow_diagram()
print(f"✓ 流程图: {flow_path}")

print("\n正在生成Word文档...")

# 创建Word文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'SimSun'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# 封面
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('Field Info Agent\n现场信息收集智能体')
title_run.font.name = 'SimHei'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
title_run.font.size = Pt(26)
title_run.font.bold = True

doc.add_paragraph()

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run('技术方案与业务场景设计书（完整版）')
subtitle_run.font.name = 'SimSun'
subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
subtitle_run.font.size = Pt(16)
subtitle_run.font.bold = True

for _ in range(4):
    doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_para.add_run('版本：V1.0\n日期：2026年3月')
info_run.font.name = 'SimSun'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
info_run.font.size = Pt(12)

doc.add_page_break()

# 插入架构图
doc.add_heading('一、 整体业务架构', level=1)
arch_para = doc.add_paragraph()
arch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = arch_para.add_run()
run.add_picture(arch_path, width=Inches(6))

caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = caption.add_run('图1-1 Field Info Agent整体业务架构')
cap_run.font.name = 'SimSun'
cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
cap_run.font.size = Pt(10)

doc.add_paragraph()
para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('Field Info Agent采用分层架构设计，分为用户交互层、渠道接入层、智能处理层、工具执行层和数据存储层五个层次。系统基于OpenClaw智能体框架构建，充分利用企业微信长连接、KIMI大模型多模态能力、PostgreSQL关系型数据库和MinIO对象存储，形成完整的技术体系。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

doc.add_page_break()

# 插入状态流转图
doc.add_heading('二、 Session状态流转', level=1)
state_para = doc.add_paragraph()
state_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = state_para.add_run()
run.add_picture(state_path, width=Inches(5.5))

caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = caption.add_run('图2-1 Session状态流转图')
cap_run.font.name = 'SimSun'
cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
cap_run.font.size = Pt(10)

doc.add_paragraph()
para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('系统采用状态机模式管理会话生命周期。从IDLE空闲状态开始，用户启动驻点后进入PREPARING准备状态，确认开始后进入COLLECTING采集状态。采集完成后进入ANALYZING分析状态进行AI批量分析，最终进入COMPLETED完成状态并生成文档。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

doc.add_page_break()

# 插入时序图
doc.add_heading('三、 核心业务时序', level=1)
flow_para = doc.add_paragraph()
flow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = flow_para.add_run()
run.add_picture(flow_path, width=Inches(6))

caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = caption.add_run('图3-1 核心业务时序图 - 照片分析流程')
cap_run.font.name = 'SimSun'
cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
cap_run.font.size = Pt(10)

doc.add_paragraph()
para = doc.add_paragraph()
para.paragraph_format.first_line_indent = Cm(0.5)
para.paragraph_format.line_spacing = 1.5
run = para.add_run('照片分析流程展示了用户、企业微信、Agent和KIMI之间的完整交互。用户发送照片后，通过WebSocket长连接实时推送到Agent，Agent调用KIMI多模态API进行分析，然后通过流式消息将结果返回给用户。整个过程采用长连接模式，无需异步队列，响应更及时。')
run.font.name = 'SimSun'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

# 保存
doc.save('Field_Info_Agent_完整方案_含图表_v1.0.docx')
print("\n✅ Word文档已生成：Field_Info_Agent_完整方案_含图表_v1.0.docx")

# 清理
shutil.rmtree("temp_images", ignore_errors=True)
print("✅ 临时文件已清理")
